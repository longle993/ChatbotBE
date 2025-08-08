from flask import Blueprint, jsonify, request
from .db import mongo
import pandas as pd
import os
from dotenv import load_dotenv
from langchain_core.documents import Document
from services.embed import Embedder
from services.vector_db import VectorDB
from services.chatbot import Gemini
import base64

chatbot_api = Blueprint('chatbot_api', __name__)

load_dotenv()
# Model nhúng
model_name = os.getenv("EMMBEDDING_MODEL")

# Khởi tạo embedder, vector_db và chatbot
embedder = Embedder(model_name)
persist_path = os.getenv("PERSIST_PATH")
vector_db = VectorDB(embedder.model, persist_path=persist_path)
gemini_chatbot = Gemini(vector_db.get_vectorstore())

@chatbot_api.route('/user', methods=['GET'])
def get_user():
    db = mongo['Chatbot']
    users = list(db.user.find())
    for user in users:
        user['_id'] = str(user['_id'])
    return jsonify({
        "code": 200,
        "message": "Success",
        "data": users
    })

@chatbot_api.route('/user', methods=['POST'])
def login_user():
    db = mongo['Chatbot']
    data = request.get_json()
    username = data.get("username")
    password = data.get("password")
    if not username or not password:
        return jsonify({
            "code": 400,
            "message": "Missing username or password"
        }), 400

    user = db.user.find_one({"username": username, "password": password})
    if user:
        user['_id'] = str(user['_id'])
        return jsonify({
            "code": 200,
            "message": "Login successful",
            "data": {"username": user['username']}
        }), 200
    else:
        return jsonify({
            "code": 401,
            "message": "Invalid username or password"
        }), 401

@chatbot_api.route('/message', methods=['POST'])
def chat_message():
    data = request.get_json()
    user_input = data.get("message")
    files = data.get("files", [])

    if not user_input and not files:
        return jsonify({
            "code": 400,
            "message": "Missing message or files"
        }), 400

    # Đọc nội dung từ files và nối vào context
    file_contents = []
    for f in files:
        text = ""
        file_name = f.get("name", "").lower()
        encoded_content = f.get("content", "")

        try:
            file_bytes = base64.b64decode(encoded_content.split(',')[-1])
        except Exception:
            file_bytes = b""

        if file_name.endswith(".pdf"):
            # Đọc PDF (cần cài PyPDF2 hoặc pdfplumber)
            try:
                import io
                from PyPDF2 import PdfReader
                pdf_reader = PdfReader(io.BytesIO(file_bytes))
                for page in pdf_reader.pages:
                    t = page.extract_text() or ""
                    text += t
            except Exception:
                text += "[Không thể đọc file PDF]"
        elif file_name.endswith(".docx"):
            # Đọc DOCX
            try:
                import io
                import docx
                doc = docx.Document(io.BytesIO(file_bytes))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            except Exception:
                text += "[Không thể đọc file DOCX]"
        elif file_name.endswith(".xlsx"):
            # Đọc Excel: kết hợp header + row
            try:    
                # Đọc Excel
                df = pd.read_excel("data/documents/data.xlsx")
                header = df.columns.tolist()

                # Chuyển Excel → Documents
                documents = []
                for index, row in df.iterrows():
                    content = embedder.combine_text_columns(row, header)
                    doc = Document(
                        page_content=content,
                        metadata={"index": index}
                    )
                    documents.append(doc)
                    print(content)

                # Nhúng & lưu vào FAISS
                vector_db.add_documents(documents)
                print("✅ Đã nhúng xong và lưu FAISS thành công.")
            except Exception:
                text += "[Không thể đọc file Excel]"
        else:
            # Nếu là text hoặc không xác định thì lấy trực tiếp
            try:
                text = file_bytes.decode("utf-8")
            except Exception:
                text = "[Không thể đọc file dạng text]"

        # Giới hạn 3000 ký tự cho mỗi file
        if text:
            file_contents.append(f"{f['name']}:\n{text[:3000]}")

    # Ghép nội dung file vào truy vấn
    if file_contents:
        user_input = (user_input or "") + "\n\nNội dung file:\n" + "\n\n".join(file_contents)

    try:
        response = gemini_chatbot.chat(user_input)
        return jsonify({
            "code": 200,
            "message": "Success",
            "data": {
                "response": response
            }
        })
    except Exception as e:
        return jsonify({
            "code": 500,
            "message": f"Error: {str(e)}"
        }), 500