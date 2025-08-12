# services/file_service.py
import io
import os
import base64
import docx
import pandas as pd
from PyPDF2 import PdfReader
from langchain_core.documents import Document
from typing import List, Dict, Any, Tuple, Optional

class FileService:
    """
    Service xử lý các loại file khác nhau
    """
    
    SUPPORTED_EXTENSIONS = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'txt': 'text/plain',
        'csv': 'text/csv'
    }
    
    def __init__(self):
        pass
    
    def decode_base64_file_content(encoded_content: str) -> Optional[bytes]:
        try:
            return base64.b64decode(encoded_content.split(',')[-1])
        except Exception:
            return None

    def extract_file_contents(self, files: List[Dict], max_content_length: int = 3000) -> Dict[str, Any]:
        """
        Trích xuất nội dung từ danh sách files
        
        Args:
            files (List[Dict]): Danh sách file info với 'name' và 'content' (base64)
            max_content_length (int): Độ dài tối đa của nội dung trả về
            
        Returns:
            Dict: Kết quả trích xuất với thông tin chi tiết
        """
        try:
            extracted_contents = []
            processing_results = []
            
            for file_info in files:
                result = self._extract_single_file_content(file_info, max_content_length)
                
                if result["success"]:
                    extracted_contents.append(result["content"])
                
                processing_results.append({
                    "filename": file_info.get("name", "unknown"),
                    "success": result["success"],
                    "content_length": len(result.get("content", "")),
                    "error": result.get("error")
                })
            
            return {
                "success": True,
                "extracted_contents": extracted_contents,
                "processing_results": processing_results,
                "total_files": len(files),
                "successful_files": sum(1 for r in processing_results if r["success"])
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error extracting file contents: {str(e)}",
                "extracted_contents": [],
                "processing_results": []
            }
    
    def _extract_single_file_content(self, file_info: Dict, max_length: int) -> Dict[str, Any]:
        """
        Trích xuất nội dung từ một file
        
        Args:
            file_info (Dict): Thông tin file
            max_length (int): Độ dài tối đa
            
        Returns:
            Dict: Kết quả trích xuất
        """
        try:
            file_name = file_info.get("name", "").lower()
            encoded_content = file_info.get("content", "")
            
            # Decode base64 content
            try:
                file_bytes = base64.b64decode(encoded_content.split(',')[-1])
            except Exception:
                return {
                    "success": False,
                    "error": "Invalid base64 content"
                }
            
            # Xử lý theo loại file
            if file_name.endswith(".pdf"):
                content = self._extract_pdf_content(file_bytes)
            elif file_name.endswith(".docx"):
                content = self._extract_docx_content(file_bytes)
            elif file_name.endswith(".xlsx"):
                content = self._extract_xlsx_content(file_bytes)
            elif file_name.endswith((".txt", ".csv")):
                content = self._extract_text_content(file_bytes)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_name}"
                }
            
            # Cắt ngắn nội dung nếu cần
            if len(content) > max_length:
                content = content[:max_length] + "... [Content truncated]"
            
            formatted_content = f"{file_info.get('name', 'Unknown file')}:\n{content}"
            
            return {
                "success": True,
                "content": formatted_content,
                "original_length": len(content)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error extracting content from {file_info.get('name', 'file')}: {str(e)}"
            }
    
    def _extract_pdf_content(self, file_bytes: bytes) -> str:
        """Trích xuất nội dung từ PDF"""
        try:
            pdf_reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            
            for page in pdf_reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
            
            return "\n\n".join(text_parts) if text_parts else "[Empty PDF or unable to extract text]"
            
        except Exception as e:
            return f"[Error reading PDF: {str(e)}]"
    
    def _extract_docx_content(self, file_bytes: bytes) -> str:
        """Trích xuất nội dung từ DOCX"""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            return "\n".join(paragraphs) if paragraphs else "[Empty DOCX document]"
            
        except Exception as e:
            return f"[Error reading DOCX: {str(e)}]"
    
    def _extract_xlsx_content(self, file_bytes: bytes) -> str:
        """Trích xuất nội dung từ XLSX"""
        try:
            df = pd.read_excel(io.BytesIO(file_bytes))
            
            if df.empty:
                return "[Empty Excel file]"
            
            # Tạo summary của Excel data
            summary_parts = [
                f"Excel file with {len(df)} rows and {len(df.columns)} columns",
                f"Columns: {', '.join(df.columns.tolist())}",
                "\nFirst 5 rows:"
            ]
            
            # Thêm 5 dòng đầu
            for idx, row in df.head().iterrows():
                row_data = []
                for col in df.columns:
                    value = row[col]
                    if pd.notna(value):
                        row_data.append(f"{col}: {value}")
                if row_data:
                    summary_parts.append(f"Row {idx + 1}: {', '.join(row_data)}")
            
            return "\n".join(summary_parts)
            
        except Exception as e:
            return f"[Error reading Excel: {str(e)}]"
    
    def _extract_text_content(self, file_bytes: bytes) -> str:
        """Trích xuất nội dung từ text files"""
        try:
            # Thử các encoding khác nhau
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return file_bytes.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            return "[Unable to decode text file with supported encodings]"
            
        except Exception as e:
            return f"[Error reading text file: {str(e)}]"
    
    def process_files_to_documents(self, files: List[Dict], embedder_service=None) -> Dict[str, Any]:
        """
        Chuyển đổi files thành Documents để embedding
        
        Args:
            files (List[Dict]): Danh sách file info
            embedder_service: EmbedderService instance
            
        Returns:
            Dict: Kết quả chuyển đổi với documents
        """
        try:
            all_documents = []
            processing_results = []
            
            for file_info in files:
                result = self._process_single_file_to_documents(file_info, embedder_service)
                
                if result["success"]:
                    all_documents.extend(result["documents"])
                
                processing_results.append({
                    "filename": file_info.get("name", "unknown"),
                    "success": result["success"],
                    "document_count": len(result.get("documents", [])),
                    "error": result.get("error")
                })
            
            return {
                "success": True,
                "documents": all_documents,
                "processing_results": processing_results,
                "total_documents": len(all_documents),
                "total_files": len(files)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error processing files to documents: {str(e)}",
                "documents": [],
                "processing_results": []
            }
    
    def _process_single_file_to_documents(self, file_info: Dict, embedder_service=None) -> Dict[str, Any]:
        """
        Chuyển đổi một file thành Documents
        
        Args:
            file_info (Dict): Thông tin file
            embedder_service: EmbedderService instance
            
        Returns:
            Dict: Kết quả chuyển đổi
        """
        try:
            file_name = file_info.get("name", "").lower()
            encoded_content = file_info.get("content", "")
            
            try:
                file_bytes = base64.b64decode(encoded_content.split(',')[-1])
            except Exception:
                return {
                    "success": False,
                    "error": "Invalid base64 content",
                    "documents": []
                }
            
            documents = []
            
            if file_name.endswith(".pdf"):
                documents = self._pdf_to_documents(file_bytes, file_name)
            elif file_name.endswith(".docx"):
                documents = self._docx_to_documents(file_bytes, file_name)
            elif file_name.endswith(".xlsx") and embedder_service:
                documents = self._xlsx_to_documents(file_bytes, file_name, embedder_service)
            elif file_name.endswith((".txt", ".csv")):
                documents = self._text_to_documents(file_bytes, file_name)
            else:
                return {
                    "success": False,
                    "error": f"Unsupported file type for document conversion: {file_name}",
                    "documents": []
                }
            
            return {
                "success": True,
                "documents": documents,
                "document_count": len(documents)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error processing {file_info.get('name', 'file')} to documents: {str(e)}",
                "documents": []
            }
    
    def _pdf_to_documents(self, file_bytes: bytes, filename: str) -> List[Document]:
        """Chuyển PDF thành Documents"""
        documents = []
        try:
            pdf_reader = PdfReader(io.BytesIO(file_bytes))
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    doc = Document(
                        page_content=text,
                        metadata={
                            "source": filename,
                            "page": page_num + 1,
                            "type": "pdf"
                        }
                    )
                    documents.append(doc)
        except Exception:
            pass
        
        return documents
    
    def _docx_to_documents(self, file_bytes: bytes, filename: str) -> List[Document]:
        """Chuyển DOCX thành Documents"""
        documents = []
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            if text_parts:
                full_text = "\n".join(text_parts)
                document = Document(
                    page_content=full_text,
                    metadata={
                        "source": filename,
                        "type": "docx",
                        "paragraphs": len(text_parts)
                    }
                )
                documents.append(document)
        except Exception:
            pass
        
        return documents

    def _xlsx_to_documents(self, file_bytes: bytes, filename: str, embedder_service) -> List[Document]:
        """Chuyển XLSX thành Documents"""
        documents = []
        try:
            df = pd.read_excel(io.BytesIO(file_bytes))
            header = df.columns.tolist()
            docs = embedder_service.embed_excel(df, header)
            documents.extend(docs)
        except Exception:
            pass

        return documents
