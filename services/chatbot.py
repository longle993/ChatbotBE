from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.prompts import ChatPromptTemplate
from langchain.schema import StrOutputParser
from langchain.memory import ConversationBufferWindowMemory  # Thay đổi này
from langchain.text_splitter import RecursiveCharacterTextSplitter
import io
import docx
import base64
import pandas as pd
from langchain_core.documents import Document
from PyPDF2 import PdfReader

class Gemini:
    def __init__(self, vector_db):
        # Khởi tạo model và embeddings
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash",
            temperature=0.7,
            max_tokens=2000  # Tăng token limit
        )
        
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001"
        )
        
        # Memory cho conversation - sử dụng Window Memory để tránh quá dài
        self.memory = ConversationBufferWindowMemory(
            memory_key="chat_history",
            return_messages=True,
            k=10  # Giữ lại 10 tin nhắn gần nhất
        )
        
        # Text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        
        # Vector store
        self.vector_store = vector_db
        
        # Lưu trữ context từ lần truy vấn trước
        self.last_context = ""
        self.last_query = ""

        # Prompt template được cải thiện cho RAG với context liên tục
        self.rag_prompt = ChatPromptTemplate.from_messages([
            ("system", """Bạn là một trợ lý AI thân thiện, thông minh và hữu ích. 

        HƯỚNG DẪN TRẢ LỜI:
        1. **Ưu tiên sử dụng thông tin trong Context** để trả lời câu hỏi
        2. **Kết hợp với lịch sử trò chuyện** để hiểu ngữ cảnh liên tục
        3. **Nếu câu hỏi hiện tại liên quan đến câu hỏi trước**, hãy:
           - Tham khảo thông tin đã thảo luận trước đó
           - Mở rộng hoặc làm rõ thêm dựa trên context mới
           - Duy trì tính liên kết trong cuộc trò chuyện
        4. **Nếu có từ như "đó", "này", "tiếp tục", "thêm nữa"**, hãy liên kết với nội dung trước đó
        5. Luôn trả lời bằng cùng ngôn ngữ với câu hỏi
        6. Nếu thiếu thông tin, yêu cầu làm rõ một cách thân thiện

        Context hiện tại:
        {context}

        Lịch sử trò chuyện (để hiểu ngữ cảnh liên tục):
        {chat_history}"""),
            ("human", "{question}")
        ])
        
        self.chain = self.rag_prompt | self.llm | StrOutputParser()

    def search_similar_documents(self, query, k=5):  # Tăng k lên 5
        """Tìm kiếm documents tương tự"""
        if not self.vector_store:
            return []
        
        # Kết hợp query hiện tại với query trước đó nếu có liên quan
        enhanced_query = self._enhance_query_with_context(query)
        
        docs = self.vector_store.similarity_search(enhanced_query, k=k)
        return docs
    
    def _enhance_query_with_context(self, current_query):
        """Cải thiện query bằng cách kết hợp với context trước đó"""
        # Từ khóa chỉ ra câu hỏi liên quan đến nội dung trước
        continuation_keywords = [
            "đó", "này", "tiếp tục", "thêm", "chi tiết", "mở rộng", 
            "giải thích", "that", "this", "continue", "more", "detail"
        ]
        
        current_lower = current_query.lower()
        
        # Nếu câu hỏi có từ khóa liên tục và có query trước đó
        if any(keyword in current_lower for keyword in continuation_keywords) and self.last_query:
            enhanced_query = f"{self.last_query} {current_query}"
            return enhanced_query
        
        return current_query
    
    def chat(self, user_input):
        context = ""
        
        if self.vector_store:
            relevant_docs = self.search_similar_documents(user_input, k=5)
            context = "\n\n".join([doc.page_content for doc in relevant_docs])
            
            # Nếu context hiện tại liên quan đến context trước, kết hợp chúng
            if self._is_related_query(user_input) and self.last_context:
                context = f"{self.last_context}\n\n--- THÔNG TIN BỔ SUNG ---\n\n{context}"
        
        # Lấy lịch sử chat
        chat_history = self.memory.chat_memory.messages
        
        # Generate response
        response = self.chain.invoke({
            "context": context,
            "chat_history": chat_history,
            "question": user_input
        })
        
        # Lưu vào memory
        self.memory.chat_memory.add_user_message(user_input)
        self.memory.chat_memory.add_ai_message(response)
        
        # Cập nhật context và query để sử dụng cho lần sau
        self.last_context = context[:2000]  # Giữ 2000 ký tự đầu
        self.last_query = user_input
        
        return response
    
    def _is_related_query(self, current_query):
        """Kiểm tra xem câu hỏi hiện tại có liên quan đến câu trước không"""
        continuation_indicators = [
            "đó", "này", "tiếp tục", "thêm", "chi tiết hơn", "mở rộng", 
            "giải thích rõ hơn", "that", "this", "continue", "more details",
            "expand", "elaborate", "what about", "còn", "thì sao"
        ]
        
        return any(indicator in current_query.lower() for indicator in continuation_indicators)
    
    def clear_context(self):
        """Xóa context để bắt đầu chủ đề mới"""
        self.last_context = ""
        self.last_query = ""
        self.memory.clear()
    
    def get_conversation_summary(self):
        """Lấy tóm tắt cuộc trò chuyện"""
        messages = self.memory.chat_memory.messages
        if not messages:
            return "Chưa có cuộc trò chuyện nào."
        
        summary = "Tóm tắt cuộc trò chuyện:\n"
        for i, msg in enumerate(messages[-6:], 1):  # Lấy 6 tin nhắn cuối
            role = "Người dùng" if msg.type == "human" else "AI"
            content = msg.content[:100] + "..." if len(msg.content) > 100 else msg.content
            summary += f"{i}. {role}: {content}\n"
        
        return summary
    
    # Các phương thức khác giữ nguyên...
    def extract_file_contents(self, files, embedder=None, vector_db=None):
        file_contents = []
        for f in files:
            text = ""
            file_name = f.get("name", "").lower()
            encoded_content = f.get("content", "")

            try:
                file_bytes = base64.b64decode(encoded_content.split(',')[-1])
            except Exception:
                file_bytes = b""

            if file_name.endswith(".pdf"):
                try:
                    pdf_reader = PdfReader(io.BytesIO(file_bytes))
                    for page in pdf_reader.pages:
                        t = page.extract_text() or ""
                        text += t
                except Exception:
                    text += "[Không thể đọc file PDF]"
            elif file_name.endswith(".docx"):
                try:
                    doc = docx.Document(io.BytesIO(file_bytes))
                    for para in doc.paragraphs:
                        text += para.text + "\n"
                except Exception:
                    text += "[Không thể đọc file DOCX]"
            elif file_name.endswith(".xlsx") and embedder and vector_db:
                try:
                    df = pd.read_excel(io.BytesIO(file_bytes))
                    header = df.columns.tolist()
                    documents = []
                    for index, row in df.iterrows():
                        content = embedder.combine_text_columns(row, header)
                        doc = Document(
                            page_content=content,
                            metadata={"index": index}   
                        )
                        documents.append(doc)
                    vector_db.add_documents(documents)
                except Exception as e:
                    text += "[Không thể đọc file Excel]" + str(e)
            else:
                try:
                    text = file_bytes.decode("utf-8")
                except Exception:
                    text = "[Không thể đọc file dạng text]"

            if text:
                file_contents.append(f"{f['name']}:\n{text[:3000]}")
        return file_contents
        
    def embed_and_reload_files(self, files, embedder, vector_db):
        """
        Nhúng nội dung các file vào vector store và load lại vector store cho chatbot.
        """
        documents = []
        for f in files:
            file_name = f.get("name", "").lower()
            encoded_content = f.get("content", "")
            try:
                file_bytes = base64.b64decode(encoded_content.split(',')[-1])
            except Exception:
                file_bytes = b""

            if file_name.endswith(".pdf"):
                try:
                    pdf_reader = PdfReader(io.BytesIO(file_bytes))
                    for page in pdf_reader.pages:
                        text = page.extract_text() or ""
                        doc = Document(page_content=text, metadata={"name": file_name})
                        documents.append(doc)
                except Exception:
                    pass
            elif file_name.endswith(".docx"):
                try:
                    docx_file = docx.Document(io.BytesIO(file_bytes))
                    text = "\n".join([para.text for para in docx_file.paragraphs])
                    doc = Document(page_content=text, metadata={"name": file_name})
                    documents.append(doc)
                except Exception:
                    pass
            elif file_name.endswith(".xlsx"):
                try:
                    df = pd.read_excel(io.BytesIO(file_bytes))
                    header = df.columns.tolist()
                    for index, row in df.iterrows():
                        content = embedder.combine_text_columns(row, header)
                        doc = Document(page_content=content, metadata={"name": file_name, "index": index})
                        documents.append(doc)
                except Exception:
                    pass
            else:
                try:
                    text = file_bytes.decode("utf-8")
                    doc = Document(page_content=text, metadata={"name": file_name})
                    documents.append(doc)
                except Exception:
                    pass

        if documents:
            vector_db.add_documents(documents)
            self.vector_store = vector_db.get_vectorstore()


# Ví dụ sử dụng:
"""
# Khởi tạo chatbot
chatbot = Gemini(vector_db)

# Cuộc trò chuyện liên tục
user1 = "FBO là gì?"
response1 = chatbot.chat(user1)
print(f"Bot: {response1}")

user2 = "Chi tiết hơn về nội dung đó"  # Bot sẽ hiểu "đó" là FBO
response2 = chatbot.chat(user2)
print(f"Bot: {response2}")

user3 = "Ví dụ thực tế thì sao?"  # Tiếp tục về chủ đề FBO
response3 = chatbot.chat(user3)
print(f"Bot: {response3}")

# Xóa context khi muốn chuyển chủ đề mới
chatbot.clear_context()
"""